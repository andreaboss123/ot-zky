#!/usr/bin/env python3
"""
Skript pro generování Word dokumentu (.docx) s maturitními otázkami z Markdown souboru.
Script for generating Word document (.docx) with maturita questions from Markdown file.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_page_margins(section, margin=2.5):
    """Nastaví okraje stránky / Set page margins"""
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)


def add_page_number(section):
    """Přidá číslování stránek / Add page numbering"""
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def set_font(run, font_name='Times New Roman', size=12, bold=False):
    """Nastaví font / Set font"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    # Nastavení pro kompatibilitu s různými verzemi Word
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def create_title_page(doc):
    """Vytvoří titulní stranu / Create title page"""
    # Titulní strana
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('\n\n\n\n\nMaturitní otázky\n')
    set_font(title_run, size=24, bold=True)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Účetnictví\n')
    set_font(subtitle_run, size=20, bold=True)
    
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2_run = subtitle2.add_run('Česká republika')
    set_font(subtitle2_run, size=16)
    
    # Konec titulní strany - nová stránka
    doc.add_page_break()


def create_table_of_contents(doc, chapters):
    """Vytvoří obsah / Create table of contents"""
    toc_title = doc.add_paragraph()
    toc_run = toc_title.add_run('Obsah')
    set_font(toc_run, size=16, bold=True)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    for i, chapter_title in enumerate(chapters, 1):
        toc_entry = doc.add_paragraph()
        toc_run = toc_entry.add_run(f'{i}. {chapter_title}')
        set_font(toc_run)
        toc_entry.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    doc.add_page_break()


def parse_markdown(md_file):
    """Parsuje Markdown soubor a vrací seznam kapitol / Parse Markdown file and return list of chapters"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Rozdělení na kapitoly podle ## nadpisů
    chapters = []
    current_chapter = None
    
    for line in content.split('\n'):
        if line.startswith('## '):
            if current_chapter:
                chapters.append(current_chapter)
            current_chapter = {
                'title': line[3:].strip(),
                'content': []
            }
        elif current_chapter is not None:
            current_chapter['content'].append(line)
    
    if current_chapter:
        chapters.append(current_chapter)
    
    return chapters


def add_chapter(doc, chapter_num, chapter):
    """Přidá kapitolu do dokumentu / Add chapter to document"""
    # STRÁNKA 1: Pouze otázka (vycentrovaná)
    # Question page - centered and prominent
    question_para = doc.add_paragraph()
    question_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Přidání vertikálního prostoru před otázkou
    for _ in range(8):
        doc.add_paragraph()
    
    # Samotná otázka
    question_heading = doc.add_paragraph()
    question_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    question_run = question_heading.add_run(f'{chapter_num}. {chapter["title"]}')
    set_font(question_run, size=16, bold=True)
    question_heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # Nová stránka pro odpověď
    doc.add_page_break()
    
    # STRÁNKA 2: Odpověď
    # Answer page with content
    # Nadpis odpovědi (menší než otázka)
    answer_heading = doc.add_paragraph()
    answer_heading_run = answer_heading.add_run(f'Odpověď {chapter_num}:')
    set_font(answer_heading_run, size=14, bold=True)
    answer_heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    doc.add_paragraph()
    
    # Obsah odpovědi
    content_text = '\n'.join(chapter['content']).strip()
    
    # Rozdělení na odstavce
    paragraphs = content_text.split('\n\n')
    
    for para_text in paragraphs:
        if para_text.strip():
            # Odstranení markdown formátování
            para_text = re.sub(r'\*\*(.*?)\*\*', r'\1', para_text)  # bold
            para_text = re.sub(r'\*(.*?)\*', r'\1', para_text)  # italic
            
            para = doc.add_paragraph()
            para_run = para.add_run(para_text.strip())
            set_font(para_run)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Nová stránka po odpovědi
    doc.add_page_break()


def generate_document(md_file, output_file):
    """Hlavní funkce pro generování dokumentu / Main function for document generation"""
    print(f'Načítám Markdown soubor: {md_file}')
    
    # Kontrola existence souboru
    if not os.path.exists(md_file):
        print(f'CHYBA: Soubor {md_file} neexistuje!')
        print('Vytvořte prosím soubor docs/src/otazky.md s obsahem otázek.')
        return False
    
    # Parsování Markdown
    chapters = parse_markdown(md_file)
    
    if not chapters:
        print('CHYBA: V Markdown souboru nebyly nalezeny žádné kapitoly!')
        print('Použijte ## pro označení kapitol, např: ## 1. Otázka')
        return False
    
    print(f'Nalezeno {len(chapters)} kapitol')
    
    # Vytvoření dokumentu
    doc = Document()
    
    # Nastavení okrajů
    for section in doc.sections:
        set_page_margins(section)
    
    # Titulní strana
    print('Generuji titulní stranu...')
    create_title_page(doc)
    
    # Obsah
    print('Generuji obsah...')
    chapter_titles = [ch['title'] for ch in chapters]
    create_table_of_contents(doc, chapter_titles)
    
    # Nastavení číslování stránek od této stránky (po obsahu)
    section = doc.sections[-1]
    add_page_number(section)
    
    # Kapitoly
    print('Generuji kapitoly...')
    for i, chapter in enumerate(chapters, 1):
        print(f'  - Kapitola {i}: {chapter["title"]}')
        add_chapter(doc, i, chapter)
    
    # Uložení dokumentu
    print(f'Ukládám dokument: {output_file}')
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    doc.save(output_file)
    
    print('✅ Dokument byl úspěšně vygenerován!')
    print(f'📄 Výstupní soubor: {output_file}')
    return True


if __name__ == '__main__':
    # Cesty k souborům
    script_dir = os.path.dirname(os.path.abspath(__file__))
    repo_root = os.path.dirname(script_dir)
    
    md_file = os.path.join(repo_root, 'docs', 'src', 'otazky.md')
    output_file = os.path.join(repo_root, 'output', 'maturitni_otazky_ucetnictvi.docx')
    
    print('=' * 60)
    print('Generátor maturitních otázek - Účetnictví (ČR)')
    print('=' * 60)
    print()
    
    success = generate_document(md_file, output_file)
    
    if not success:
        exit(1)
